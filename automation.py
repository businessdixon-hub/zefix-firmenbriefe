"""
Zefix Firmenbrief-Automation
Täglich: neue GmbH/AG in der Schweiz → entweder Direkt-Email ODER PDF-Brief.
"""

import os
import re
import sys
import shutil
import smtplib
import logging
import datetime
import subprocess
from pathlib import Path
from email import encoders
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

import requests
from docx import Document

import email_finder
import email_template

# ── Konfiguration ──────────────────────────────────────────────────────────────
TEMPLATE_PATH   = Path("Brief an neue Firmen.docx")
PROCESSED_FILE  = Path(".processed_firms.txt")
EMAILED_FILE    = Path(".emailed_firms.txt")
LOG_FILE        = Path(".task_log.txt")
TEMP_DIR        = Path("/tmp/zefix_briefe")
OUTPUT_DIR      = Path("pdfs")
MAX_FIRMS       = int(os.environ.get("MAX_FIRMS", "10"))

# Sammel-Email an Confidio (mit PDFs für Postversand)
EMAIL_TO        = os.environ.get("EMAIL_TO", "fabio.dixon@confidio.ch")
GMAIL_USER      = os.environ.get("GMAIL_USER", "")
GMAIL_PASSWORD  = os.environ.get("GMAIL_APP_PASSWORD", "")

# Direkt-Versand an Firmen via Microsoft 365 (info@confidio.ch)
M365_USER       = os.environ.get("M365_USER", "")
M365_PASSWORD   = os.environ.get("M365_PASSWORD", "")
M365_HOST       = os.environ.get("M365_HOST", "smtp.office365.com")
M365_PORT       = int(os.environ.get("M365_PORT", "587"))

# Test-Modus: Direkt-Emails gehen statt an Firma an diese Adresse
TEST_MODE       = os.environ.get("TEST_MODE", "true").lower() == "true"
TEST_REDIRECT   = os.environ.get("TEST_REDIRECT", "fabio.dixon@confidio.ch")

SHAB_API     = "https://www.shab.ch/api/v1/publications"
ZEFIX_SEARCH = "https://www.zefix.ch/ZefixREST/api/v1/firm/search.json"
ZEFIX_DETAIL = "https://www.zefix.ch/ZefixREST/api/v1/firm/{ehraid}.json"

# ── Logging ────────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger(__name__)


# ══════════════════════════════════════════════════════════════════════════════
# 1. Tracking
# ══════════════════════════════════════════════════════════════════════════════
def _load_set(path: Path) -> set:
    if path.exists():
        return {l.strip() for l in path.read_text("utf-8").splitlines() if l.strip()}
    return set()


def _append(path: Path, line: str):
    with path.open("a", encoding="utf-8") as f:
        f.write(line + "\n")


def load_processed() -> set:
    return _load_set(PROCESSED_FILE)


def load_emailed() -> set:
    return _load_set(EMAILED_FILE)


def mark_processed(name: str):
    _append(PROCESSED_FILE, name)


def mark_emailed(name: str, email: str):
    _append(EMAILED_FILE, f"{name}\t{email}")


# ══════════════════════════════════════════════════════════════════════════════
# 2. Neue Firmen via SHAB API
# ══════════════════════════════════════════════════════════════════════════════
def get_new_gmbh_ag_this_month() -> list[dict]:
    today = datetime.date.today()
    first = today.replace(day=1)
    results = []

    for page in range(20):
        url = (
            f"{SHAB_API}?publicationStates=PUBLISHED&rubricIds=HR"
            f"&dateFrom={first}&dateTo={today}"
            f"&pageRequest.size=100&pageRequest.page={page}"
        )
        try:
            r = requests.get(url, timeout=20)
            r.raise_for_status()
        except Exception as e:
            log.error(f"SHAB API Fehler Seite {page}: {e}")
            break

        pubs = r.json().get("content", [])
        if not pubs:
            break

        for pub in pubs:
            meta = pub["meta"]
            if meta["subRubric"] != "HR01":
                continue

            title_de = meta["title"].get("de", "")
            has_gmbh = "GmbH" in title_de or "GMBH" in title_de
            has_ag   = bool(re.search(r'\bAG\b', title_de))
            if not (has_gmbh or has_ag):
                continue

            name = title_de.replace("Neueintragung ", "").rsplit(",", 1)[0].strip()
            if not name:
                continue

            results.append({"name": name, "pub_date": meta["publicationDate"][:10]})

    seen = set()
    unique = []
    for r in results:
        if r["name"] not in seen:
            seen.add(r["name"])
            unique.append(r)

    log.info(f"SHAB: {len(unique)} GmbH/AG-Neueintragungen im Monat (dedupliziert)")
    return unique


# ══════════════════════════════════════════════════════════════════════════════
# 3. Firmendaten via Zefix API
# ══════════════════════════════════════════════════════════════════════════════
def get_zefix_details(firm_name: str) -> dict | None:
    for search_type in ("EXACT", "STARTSWITH"):
        try:
            r = requests.post(ZEFIX_SEARCH, json={
                "languageKey": "de", "maxEntries": 5, "name": firm_name,
                "searchType": search_type, "status": "ACTIVE", "deletedFirms": False,
            }, timeout=20)
            firms = r.json().get("list", [])
            if not firms:
                continue
            firm = sorted(firms, key=lambda f: f.get("shabDate", ""), reverse=True)[0]
            ehraid = firm["ehraid"]
            return requests.get(ZEFIX_DETAIL.format(ehraid=ehraid), timeout=20).json()
        except Exception as e:
            log.warning(f"Zefix Fehler für '{firm_name}' ({search_type}): {e}")
    return None


def extract_address(detail: dict) -> tuple[str, str]:
    addr = detail.get("address", {}) or {}
    strasse = f"{addr.get('street','')} {addr.get('houseNumber','')}".strip() or "—"
    plz_ort = f"{addr.get('swissZipCode','')} {addr.get('town','')}".strip() or "—"
    return strasse, plz_ort


def extract_contact_person(detail: dict) -> tuple[str, str, str]:
    roles = detail.get("roles") or []
    for quality in ("mit Einzelunterschrift", "mit Kollektivunterschrift"):
        for role in roles:
            if quality in (role.get("signatureQuality") or ""):
                fn = (role.get("firstName") or "").strip()
                ln = (role.get("lastName")  or "").strip()
                if fn and ln:
                    return fn, ln, _guess_gender(fn)

    msg = (detail.get("shabPub") or [{}])[0].get("message", "") or ""
    idx = msg.find("Eingetragene Personen:")
    if idx >= 0:
        persons_block = msg[idx:]
        pattern_einzel = re.compile(
            r"([A-ZÄÖÜ][a-zäöüß\-]+),\s+"
            r"([A-ZÄÖÜ][a-zäöüßÄÖÜ\s\-]+?),\s+"
            r"von\s+.+?,\s+in\s+.+?,\s+[^,]+,\s+mit Einzelunterschrift"
        )
        pattern_kollektiv = re.compile(
            r"([A-ZÄÖÜ][a-zäöüß\-]+),\s+"
            r"([A-ZÄÖÜ][a-zäöüßÄÖÜ\s\-]+?),\s+"
            r"von\s+.+?,\s+in\s+.+?,\s+(?:Geschäftsführer|Verwaltungsrat)[^,]*,\s+mit Kollektivunterschrift"
        )
        for pattern in (pattern_einzel, pattern_kollektiv):
            m = pattern.search(persons_block)
            if m:
                nachname = m.group(1).strip()
                vorname  = m.group(2).strip().split()[0]
                return vorname, nachname, _guess_gender(vorname)

    return "", "Geschäftsführung", ""


def _guess_gender(vorname: str) -> str:
    weiblich_endungen = ("a", "e", "i", "ie", "ine", "tte", "lle", "nne", "ise")
    return "Frau" if any(vorname.lower().endswith(e) for e in weiblich_endungen) else "Herr"


# ══════════════════════════════════════════════════════════════════════════════
# 4. DOCX befüllen
# ══════════════════════════════════════════════════════════════════════════════
def _replace_in_para(para, old: str, new: str):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for run in para.runs[1:]:
            run.text = ""


def create_personalized_docx(firma, kontakt, strasse, plz_ort, anrede_formal, monat, out_path):
    shutil.copy2(TEMPLATE_PATH, out_path)
    doc = Document(out_path)

    monat_patterns = [
        f"{m} {y}"
        for m in ("Januar","Februar","März","April","Mai","Juni",
                  "Juli","August","September","Oktober","November","Dezember")
        for y in range(2024, 2030)
    ]
    replacements = {
        "Firma":                   firma,
        "Anrede Vorname Nachname": kontakt,
        "Adresse":                 strasse,
        "Postleitzahl Ort":        plz_ort,
        "Formelle Anrede":         anrede_formal,
    }

    for para in doc.paragraphs:
        for old, new in replacements.items():
            _replace_in_para(para, old, new)
        for mp in monat_patterns:
            if mp in "".join(r.text for r in para.runs):
                _replace_in_para(para, mp, monat)
                break

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        _replace_in_para(para, old, new)

    doc.save(out_path)


# ══════════════════════════════════════════════════════════════════════════════
# 5. PDF-Konvertierung
# ══════════════════════════════════════════════════════════════════════════════
def convert_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(out_dir), str(docx_path)
    ], check=True, capture_output=True)
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF nicht erstellt: {pdf_path}")
    return pdf_path


# ══════════════════════════════════════════════════════════════════════════════
# 6a. Direkt-Email an Firma (Microsoft 365)
# ══════════════════════════════════════════════════════════════════════════════
def send_direct_email(to_email: str, firm_name: str, anrede_formal: str) -> bool:
    """Sendet HTML-Email direkt an die Firma. Returns True bei Erfolg."""
    if not (M365_USER and M365_PASSWORD):
        log.warning("M365_USER/M365_PASSWORD nicht gesetzt – kann keine Direkt-Email senden")
        return False

    subject = email_template.SUBJECT
    actual_to = to_email
    if TEST_MODE:
        actual_to = TEST_REDIRECT
        subject = f"[TEST → {to_email}] {subject}"

    msg = MIMEMultipart("alternative")
    msg["From"]    = M365_USER
    msg["To"]      = actual_to
    msg["Subject"] = subject
    msg["Reply-To"] = M365_USER

    text_body = email_template.build_text(anrede_formal)
    html_body = email_template.build_html(anrede_formal, firm_name)
    msg.attach(MIMEText(text_body, "plain", "utf-8"))
    msg.attach(MIMEText(html_body, "html", "utf-8"))

    try:
        with smtplib.SMTP(M365_HOST, M365_PORT, timeout=30) as server:
            server.ehlo()
            server.starttls()
            server.ehlo()
            server.login(M365_USER, M365_PASSWORD)
            server.send_message(msg)
        log.info(f"  ✓ Direkt-Email an {actual_to} gesendet ({'TEST' if TEST_MODE else 'LIVE'})")
        return True
    except Exception as e:
        log.error(f"  ✗ Direkt-Email an {actual_to} fehlgeschlagen: {e}")
        return False


# ══════════════════════════════════════════════════════════════════════════════
# 6b. Sammel-Email an Confidio (PDFs für Postversand + Bericht)
# ══════════════════════════════════════════════════════════════════════════════
def send_summary_email(pdf_paths: list[Path], emailed: list[tuple[str, str]],
                       skipped: list[str], monat: str):
    """
    pdf_paths: PDFs für Postversand (Firmen ohne Email)
    emailed:   [(firmname, email), ...]
    skipped:   Firmen ohne Email UND ohne PDF (Fehler)
    """
    n_pdfs    = len(pdf_paths)
    n_emailed = len(emailed)
    n_skipped = len(skipped)

    if n_pdfs == 0 and n_emailed == 0 and n_skipped == 0:
        log.info("Nichts zu berichten – keine Sammel-Email")
        return

    msg = MIMEMultipart()
    msg["From"]    = GMAIL_USER
    msg["To"]      = EMAIL_TO
    mode = "TEST" if TEST_MODE else "LIVE"
    msg["Subject"] = f"[{mode}] Firmenbriefe {monat} – {n_emailed} Email / {n_pdfs} Brief"

    body_lines = [
        f"Bericht für {monat}",
        "",
        f"  Direkt per Email kontaktiert: {n_emailed}",
        f"  PDFs für Postversand:         {n_pdfs}",
    ]
    if n_skipped:
        body_lines.append(f"  Übersprungen (Fehler):       {n_skipped}")
    body_lines.append("")

    if emailed:
        body_lines.append("── Direkt per Email kontaktiert ──")
        for name, em in emailed:
            body_lines.append(f"  • {name}  →  {em}")
        body_lines.append("")

    if pdf_paths:
        body_lines.append("── PDFs für Postversand (anbei) ──")
        for p in pdf_paths:
            body_lines.append(f"  • {p.name}")
        body_lines.append("")

    if skipped:
        body_lines.append("── Übersprungen ──")
        for s in skipped:
            body_lines.append(f"  • {s}")
        body_lines.append("")

    if TEST_MODE:
        body_lines.append("⚠ TEST-MODUS: Direkt-Emails wurden NICHT an die Firmen gesendet,")
        body_lines.append(f"  sondern an {TEST_REDIRECT} umgeleitet.")

    body_lines.append("")
    body_lines.append("Diese E-Mail wurde automatisch generiert.")
    msg.attach(MIMEText("\n".join(body_lines), "plain", "utf-8"))

    for pdf in pdf_paths:
        with pdf.open("rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{pdf.name}"')
        msg.attach(part)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(GMAIL_USER, GMAIL_PASSWORD)
        server.send_message(msg)

    log.info(f"Sammel-Email gesendet: {n_emailed} Email + {n_pdfs} PDFs an {EMAIL_TO}")


# ══════════════════════════════════════════════════════════════════════════════
# 7. Hauptprogramm
# ══════════════════════════════════════════════════════════════════════════════
def main():
    today = datetime.date.today()
    monat_namen = ["Januar","Februar","März","April","Mai","Juni",
                   "Juli","August","September","Oktober","November","Dezember"]
    monat = f"{monat_namen[today.month - 1]} {today.year}"

    log.info(f"=== Start: {today} / {monat} | TEST_MODE={TEST_MODE} ===")

    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    processed = load_processed()
    log.info(f"Bereits verarbeitet: {len(processed)} Firmen")

    candidates = get_new_gmbh_ag_this_month()
    new_firms = [f for f in candidates if f["name"] not in processed]
    log.info(f"Neue (noch nicht verarbeitet): {len(new_firms)}")

    to_process = new_firms[:MAX_FIRMS]
    created_pdfs: list[Path] = []
    emailed: list[tuple[str, str]] = []
    skipped: list[str] = []

    for entry in to_process:
        firma = entry["name"]
        log.info(f"Verarbeite: {firma}")

        try:
            detail = get_zefix_details(firma)
            if not detail:
                log.warning(f"  Zefix: keine Daten – überspringe")
                skipped.append(f"{firma} (kein Zefix-Eintrag)")
                continue

            strasse, plz_ort  = extract_address(detail)
            vorname, nachname, anrede = extract_contact_person(detail)

            if nachname == "Geschäftsführung":
                kontakt       = "Geschäftsführung"
                anrede_formal = "Sehr geehrte Damen und Herren"
            else:
                kontakt = f"{anrede} {vorname} {nachname}".strip()
                anrede_formal = (
                    f"Sehr geehrter Herr {nachname}" if anrede == "Herr"
                    else f"Sehr geehrte Frau {nachname}"
                )

            log.info(f"  Kontakt: {kontakt} | {strasse}, {plz_ort}")

            # ── 1. Email-Suche ───────────────────────────────────────────────
            email_result = email_finder.find_email(firma, vorname, nachname)

            if email_result:
                email_addr, source = email_result
                if send_direct_email(email_addr, firma, anrede_formal):
                    emailed.append((firma, email_addr))
                    mark_emailed(firma, email_addr)
                    mark_processed(firma)
                    continue
                else:
                    log.warning(f"  Email-Versand fehlgeschlagen, fallback auf PDF")

            # ── 2. Fallback: PDF erstellen ───────────────────────────────────
            safe = re.sub(r'[\\/:*?"<>| ]', '_', firma)
            temp_docx = TEMP_DIR / f"{safe}.docx"
            create_personalized_docx(firma, kontakt, strasse, plz_ort, anrede_formal, monat, temp_docx)
            pdf_path  = convert_to_pdf(temp_docx, OUTPUT_DIR)
            created_pdfs.append(pdf_path)
            mark_processed(firma)
            log.info(f"  PDF: {pdf_path.name}")

        except Exception as e:
            log.error(f"  Fehler bei '{firma}': {e}", exc_info=True)
            skipped.append(f"{firma} (Fehler: {e})")
            continue

    # ── Sammel-Email ─────────────────────────────────────────────────────────
    if created_pdfs or emailed or skipped:
        send_summary_email(created_pdfs, emailed, skipped, monat)
    else:
        log.info("Keine neuen Firmen heute.")

    # ── Log ──────────────────────────────────────────────────────────────────
    with LOG_FILE.open("a", encoding="utf-8") as f:
        f.write(f"{datetime.datetime.now().isoformat()}: "
                f"{len(emailed)} Email + {len(created_pdfs)} PDF – {monat}\n")

    log.info(f"=== Ende: {len(emailed)} Email + {len(created_pdfs)} PDF erstellt ===")


if __name__ == "__main__":
    main()
